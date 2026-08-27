#!/usr/bin/env python3
"""Create the illustrated DOCX reproducibility guide."""
from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "data"
DOCS = ROOT / "docs"
OUT = DOCS / "Climate_Housing_Exposure_Index_Dashboard_Reproducibility_Guide.docx"
SCREENSHOT = ROOT / "screenshots" / "dashboard_overview.png"
REPORT_SCREENSHOT = ROOT / "screenshots" / "commissioner_precinct_decision_brief.png"

NAVY = "17324D"
DEEP = "0E2438"
TEAL = "168C95"
AQUA = "4CB8BE"
GOLD = "F2B134"
LIGHT = "F4F7F9"
MID = "D9E4E9"
TEXT = "273746"
MUTED = "617382"
RED = "C9463D"
WHITE = "FFFFFF"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs.get(edge)
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key in ["val", "sz", "space", "color"]:
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Liberation Sans"
    run.font.size = Pt(9)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    """Keep compact key-value rows together during pagination."""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_hyperlink(paragraph, text: str, url: str, color=TEAL) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_title_band(doc: Document, eyebrow: str, title: str, subtitle: str | None = None) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(7.0)
    shade(cell, NAVY)
    set_cell_margins(cell, top=150, bottom=150, start=190, end=190)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(eyebrow.upper())
    r.font.name = "Liberation Sans"
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(AQUA)
    r.font.all_caps = True
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r = p2.add_run(title)
    r.font.name = "Liberation Sans"
    r.font.size = Pt(21)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(WHITE)
    if subtitle:
        p3 = cell.add_paragraph()
        p3.paragraph_format.space_after = Pt(0)
        r = p3.add_run(subtitle)
        r.font.name = "Liberation Sans"
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor.from_string("D9E9ED")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_section_heading(doc: Document, number: str, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    p.style = doc.styles["Heading 1"]
    r = p.add_run(f"{number}  {title}")
    r.font.name = "Liberation Sans"
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)


def add_subheading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.font.name = "Liberation Sans"
    r.font.size = Pt(12.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)


def add_body(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(NAVY)
        r2 = p.add_run(text[len(bold_lead):])
        for rr in (r, r2):
            rr.font.name = "Liberation Sans"
            rr.font.size = Pt(10.2)
    else:
        r = p.add_run(text)
        r.font.name = "Liberation Sans"
        r.font.size = Pt(10.2)
        r.font.color.rgb = RGBColor.from_string(TEXT)


def add_step(doc: Document, number: int, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(0.48)
    table.columns[1].width = Inches(6.35)
    ncell, cell = table.rows[0].cells
    ncell.width = Inches(0.48)
    shade(ncell, TEAL)
    ncell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    np = ncell.paragraphs[0]
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = np.add_run(str(number))
    rr.font.name = "Liberation Sans"
    rr.font.size = Pt(12)
    rr.font.bold = True
    rr.font.color.rgb = RGBColor.from_string(WHITE)
    shade(cell, LIGHT)
    set_cell_margins(cell, top=85, bottom=85, start=130, end=130)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.font.name = "Liberation Sans"
    r.font.size = Pt(10.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run(text)
    r.font.name = "Liberation Sans"
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, "F1F4F6")
    set_cell_border(cell, left={"val": "single", "sz": "14", "color": TEAL})
    set_cell_margins(cell, top=100, bottom=100, start=150, end=130)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(code.strip("\n").splitlines()):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = "Liberation Mono"
        r.font.size = Pt(8.6)
        r.font.color.rgb = RGBColor.from_string(TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_key_value_table(doc: Document, rows: list[tuple[str, str]], widths=(2.15, 4.65)) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for key, val in rows:
        row = table.add_row()
        set_row_cant_split(row)
        cells = row.cells
        cells[0].width, cells[1].width = Inches(widths[0]), Inches(widths[1])
        shade(cells[0], NAVY)
        shade(cells[1], LIGHT)
        for c in cells:
            set_cell_margins(c, top=52, bottom=52, start=100, end=100)
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(key)
        r.font.name = "Liberation Sans"
        r.font.size = Pt(9.0)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(val)
        r.font.name = "Liberation Sans"
        r.font.size = Pt(9.0)
        r.font.color.rgb = RGBColor.from_string(TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def new_page(doc: Document, eyebrow: str, title: str, subtitle: str | None = None) -> None:
    # Page-break-before is robust when the preceding page happens to be full;
    # unlike an explicit break run, it does not create a blank intermediary page.
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    r = p.add_run("")
    r.font.size = Pt(1)
    add_title_band(doc, eyebrow, title, subtitle)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    styles = doc.styles
    styles["Normal"].font.name = "Liberation Sans"
    styles["Normal"].font.size = Pt(10.2)
    styles["Normal"].font.color.rgb = RGBColor.from_string(TEXT)
    styles["Heading 1"].font.name = "Liberation Sans"
    styles["Heading 2"].font.name = "Liberation Sans"

    header = section.header
    p = header.paragraphs[0]
    p.text = "CLIMATE HOUSING EXPOSURE INDEX  /  HARRIS COUNTY, TEXAS"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.runs[0]
    r.font.name = "Liberation Sans"
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(MUTED)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def build() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    report = json.loads((DATA / "data_quality_report.json").read_text())
    inventory = pd.read_csv(DATA / "gdb_layer_inventory.csv")
    counts = report["counts"]
    totals = report["totals"]

    doc = Document()
    configure_document(doc)

    # Cover page.
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, DEEP)
    set_cell_margins(cell, top=420, bottom=420, start=360, end=360)
    p = cell.paragraphs[0]
    r = p.add_run("HARRIS COUNTY, TEXAS")
    r.font.name = "Liberation Sans"
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(AQUA)
    r.font.all_caps = True
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Climate Housing Exposure Index Dashboard")
    r.font.name = "Liberation Sans"
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(WHITE)
    p = cell.add_paragraph()
    r = p.add_run("Design, deployment, and reproducibility guide")
    r.font.name = "Liberation Sans"
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string("D9E9ED")
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    r = p.add_run("Prepared package  •  August 2026")
    r.font.name = "Liberation Sans"
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GOLD)

    doc.add_paragraph()
    add_key_value_table(doc, [
        ("Standalone access", "index.html (data-self-contained GitHub Pages entry; identical compatibility HTML included)"),
        ("Full-data application", "python server.py  →  http://127.0.0.1:8050"),
        ("Rebuild scripts", "preprocess_data.py and build_dashboard.py"),
        ("Contact", "Kaifa Lu / CECREH / Kaifa.Lu@ttu.edu"),
    ])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Research  •  Education  •  Planning exploration")
    r.font.name = "Liberation Sans"
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)

    # Page 2.
    new_page(doc, "01  PRODUCT OVERVIEW", "Dashboard interface and access modes", "A county-branded interactive explorer modeled on the Harris County HCD mapping-page hierarchy")
    if SCREENSHOT.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        shape = p.add_run().add_picture(str(SCREENSHOT), width=Inches(7.05))
        shape._inline.docPr.set("title", "Climate Housing Exposure Index dashboard overview")
        shape._inline.docPr.set("descr", "Dashboard screenshot showing the Harris County map, layer controls, tract, ZIP-code and commissioner-precinct locators, Quick Decision View, key indicators, legends, and selected-tract charts.")
        cap = doc.add_paragraph("Figure 1. Data-self-contained dashboard overview with analytical layers, key indicators, and location controls.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.name = "Liberation Sans"
        cap.runs[0].font.size = Pt(8.5)
        cap.runs[0].font.italic = True
        cap.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    # Page 3.
    new_page(doc, "01B  PRODUCT OVERVIEW", "Dashboard purpose and access", "Who the dashboard serves, what it supports, and how to open or publish it")
    add_key_value_table(doc, [
        ("Dashboard overview", "An interactive Harris County platform showing how future precipitation extremes intersect with housing, population, employment, social vulnerability, and land-use change."),
        ("Key functions", "Explore four GMT warming thresholds, CHEI patterns, 2020–2050 growth, housing stocks, parcel land use, pattern-based hotspots, clickable legend filters, tract/ZIP/precinct location, four Quick Decision questions, and one-page decision briefs."),
        ("Potential audiences", "Researchers, planners, local governments, housing and community-development agencies, emergency managers, policymakers, nonprofit organizations, and community stakeholders."),
        ("Practical applications", "Climate adaptation, housing resilience, land-use and growth management, infrastructure investment, vulnerability assessment, and environmental-justice research."),
    ])
    add_subheading(doc, "Three access modes")
    add_key_value_table(doc, [
        ("GitHub Pages / standalone", "Publish index.html and .nojekyll from the repository root, or open index.html directly. Analytical layers, ZIP and precinct boundaries, legend filters, Quick Decision actions, report generation, patterns, and JavaScript are embedded; external internet access is used only for optional basemap tiles."),
        ("Full-data local app", "Run server.py. The map requests all relevant single- and multi-family locations for the current viewport."),
        ("Container deployment", "Build the included Dockerfile on a compatible host, then use the provider-issued HTTPS address as the public dashboard URL."),
    ])
    add_body(doc, "The dashboard is designed for screening and exploration. Modeled and projected values should be validated with authoritative local data before decisions are made.")

    # Page 4.
    new_page(doc, "02  DATA AUDIT", "Uploaded geodatabase and analytical coverage", "Automated inventory, consolidated tract metrics, and documented substitutions")
    add_section_heading(doc, "2.1", "Verified source scale")
    metrics = doc.add_table(rows=2, cols=4)
    metrics.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = [
        (f"{counts['census_tracts']:,}", "census tracts", TEAL),
        (f"{counts['parcels']:,}", "changed parcels", AQUA),
        (f"{counts['single_family_points']:,}", "single-family records", GOLD),
        (f"{counts['multi_family_points']:,}", "multi-family records", RED),
    ]
    for j, (value, label, color) in enumerate(labels):
        c = metrics.cell(0, j)
        shade(c, color)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        r.font.name = "Liberation Sans"
        r.font.size = Pt(17)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        c2 = metrics.cell(1, j)
        shade(c2, LIGHT)
        p = c2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.font.name = "Liberation Sans"
        r.font.size = Pt(8.8)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(NAVY)
    add_key_value_table(doc, [
        ("ZIP-code navigation", f"{counts['zip_codes']:,} Harris County ZIP boundaries are embedded for five-digit search and an optional all-boundaries overlay."),
        ("Commissioner precinct navigation", f"{counts['commissioner_precincts']:,} Harris County commissioner precinct boundaries are embedded for selection, zoom, highlighting, and an optional all-boundaries overlay."),
        ("Reporting geography", "ZIP and precinct decision briefs summarize intersecting census tracts; no official ZIP- or precinct-level source aggregation is created."),
    ])
    add_subheading(doc, "Layer groups")
    add_key_value_table(doc, [
        ("Climate precipitation", "Four GMT point layers, four tract aggregations, one GMT 2.5°C-vs-1.5°C change layer, and four reproducibly derived display surfaces."),
        ("Parcel housing and land use", "2020 current land use, 2050 projected land use, parcel housing-unit change, and single-/multi-family locations."),
        ("Tract housing and growth", "CHEI, SVI, population density, population/household/employment projections and changes, housing-stock counts, and compound hotspots."),
        ("Reference geometry", "Harris County, census-tract, parcel, 155 ZIP-code, and four commissioner-precinct boundaries prepared for web mapping. ZIP and precinct geometry support location, orientation, and tract-based screening summaries only."),
    ])
    add_subheading(doc, "Audit findings that affect the interface")
    add_body(doc, "1. The open-source FileGDB reader exposes 32 vector layers, including Harris_County_Zipcodes and Harris_County_Commissioner_Precincts, but does not expose the four user-listed gmt_*_pr_Kriging datasets. The package retains every exposed point and tract precipitation layer and produces separately labeled derived surfaces.")
    add_body(doc, "2. harris_census_tract_CHEI_2050 and harris_census_tract_climate_housing_exposure_index_2050 contain identical CHEI 2050 values (maximum absolute difference = 0) and are consolidated in the map menu.")
    add_body(doc, "3. The actual GDB names use extreme_precip; the supplied inventory text used extreme_precipi for several tract layer names.")

    # Page 4.
    new_page(doc, "03  SYSTEM DESIGN", "Data pipeline, kriging, performance, and privacy", "A static analytical bundle with an optional full-point query service")
    add_section_heading(doc, "3.1", "Processing architecture")
    for n, (title, text) in enumerate([
        ("Read and audit", "Enumerate every exposed FileGDB layer, record geometry/counts, and identify named layers that are absent or unsupported."),
        ("Consolidate tracts", "Join CHEI, SVI, GMT precipitation, growth, housing, and hotspot attributes by 11-digit GEOID."),
        ("Prepare web geometry", "Simplify tract, ZIP, commissioner-precinct, parcel, and county geometry without changing the source geodatabase; preserve parcel-level housing-unit change, land-use labels, five-digit ZIP identifiers, and stable precinct numbers."),
        ("Optimize housing stocks", "Create a complete 1 km density grid, a standalone sample, and x-sorted NumPy arrays used by the viewport API."),
        ("Build the interface", "Serialize the Bokeh application, analytical data, ZIP and precinct geometry, clickable legends, Quick Decision actions, browser report generator, hotspot patterns, and inline resources into one portable HTML file; serve the same file through FastAPI for full-point mode."),
    ], 1):
        add_step(doc, n, title, text)
    add_subheading(doc, "Derived ordinary-kriging surfaces")
    add_body(doc, "For each GMT threshold, the script fits an exponential semivariogram to the 246 uploaded climate-model points, solves an ordinary-kriging system, predicts a regular Harris County grid, masks cells outside the county, and stores both numeric and RGBA arrays. The surfaces are display products, not replacements for authoritative rasters.")
    add_subheading(doc, "Performance and privacy controls")
    add_body(doc, "The server queries x-sorted point arrays by binary search, filters by the current y-range, and deterministically caps the browser response. The exported web assets omit property account IDs, mailing fields, assessed values, and unused HCAD attributes. The separate ZIP and commissioner-precinct layers retain only fields needed for location, area context, tract-intersection screening, and generalized geometry.")

    # Page 5.
    new_page(doc, "04  INSTALLATION", "Run the prepared dashboard locally", "Two commands are sufficient after installing the pinned environment")
    add_section_heading(doc, "4.1", "Python virtual environment")
    add_code(doc, """
python -m venv .venv
# Windows: .venv\\Scripts\\activate
# macOS/Linux: source .venv/bin/activate
pip install -r requirements.txt
python server.py
""")
    add_body(doc, "Open http://127.0.0.1:8050. The health endpoint is http://127.0.0.1:8050/api/health.")
    add_section_heading(doc, "4.2", "Conda environment")
    add_code(doc, """
conda env create -f environment.yml
conda activate climate-housing-dashboard
python server.py
""")
    add_section_heading(doc, "4.3", "No-server preview")
    add_body(doc, "Open index.html directly in a modern browser. The compatibility file climate_housing_exposure_index_dashboard.html contains identical content. Tract, parcel, climate-point, grid, pattern-hotspot, ZIP and precinct locators, clickable legend filters, Quick Decision actions, one-page reports, charts, methods, and terms content remains interactive. Only the million-record single-family point service changes to the embedded deterministic preview sample. Internet access is required only for the optional basemap tiles.")
    add_section_heading(doc, "4.4", "Package structure")
    add_code(doc, """
Climate_Housing_Exposure_Index_Dashboard/
├── index.html / .nojekyll
├── climate_housing_exposure_index_dashboard.html
├── server.py
├── preprocess_data.py
├── build_dashboard.py
├── data/
├── docs/
├── screenshots/
├── scripts/validate_dashboard.py
├── requirements.txt / environment.yml
└── Dockerfile / render.yaml
""")

    # Page 6.
    new_page(doc, "05  FULL REBUILD", "Regenerate the dashboard from the File Geodatabase", "The source geodatabase is read-only throughout this workflow")
    add_step(doc, 1, "Extract the uploaded archive", "Unzip Climate_Housing_Exposure_Index_Dashboard.gdb_boundary.zip and identify the extracted Climate_Housing_Exposure_Index_Dashboard.gdb directory.")
    add_step(doc, 2, "Run preprocessing", "Pass the extracted geodatabase path and the package data directory to preprocess_data.py.")
    add_code(doc, """
python preprocess_data.py \\
  --gdb "/path/to/Climate_Housing_Exposure_Index_Dashboard.gdb" \\
  --output data
""")
    add_step(doc, 3, "Rebuild the standalone application", "Serialize the prepared data, charts, controls, methods, source links, and terms into the HTML deliverable.")
    add_code(doc, "python build_dashboard.py")
    add_step(doc, 4, "Validate the result", "Check required files, the 32-layer inventory, 155 ZIP boundaries, four commissioner precincts, sorted point arrays, hotspot logic, pattern identifiers, clickable legends, decision tools, locator controls, and source-contact text.")
    add_code(doc, "python scripts/validate_dashboard.py")
    add_subheading(doc, "Principal generated assets")
    add_key_value_table(doc, [
        ("tracts_web.geojson", "1,115 consolidated tract features with all map, profile, legend-filter, and report fields"),
        ("zipcodes_web.geojson / commissioner_precincts_web.geojson", "155 ZIP boundaries and four precinct boundaries for search, optional overlays, and tract-based screening reports"),
        ("parcels_web.geojson / housing_grid.geojson", "20,344 changed parcels and 4,863 occupied 1 km housing-density cells"),
        ("Point arrays, kriging assets, and data-quality report", "Full x-sorted housing arrays, derived precipitation grids and model metadata, counts, audit notes, and privacy documentation"),
    ])

    # Page 7.
    new_page(doc, "06  DASHBOARD USE", "Layer exploration, location, and decision support", "Controls are intentionally transparent and interpretable")
    add_section_heading(doc, "6.1", "Primary map layer")
    add_body(doc, "Choose a composite, climate, growth, equity, housing, or parcel layer. The legend, histogram, hover value, opacity, and about-this-layer panel update together. Click one legend range or category to highlight matching features, fade nonmatches, view a live count, zoom to matches, or clear the filter. A GMT selector appears only for tract precipitation, climate points, and kriging surfaces. A housing-type selector appears only for density and point views.")
    add_section_heading(doc, "6.2", "Compound-hotspot typology")
    add_body(doc, "The noncompensatory hotspot layer retains all eight combinations of high precipitation hazard, projected household growth, and social vulnerability: none; hazard only; growth only; SVI only; hazard plus growth; hazard plus SVI; growth plus SVI; and all three. Patterns, rather than color alone, encode the factors: diagonal lines indicate precipitation hazard, vertical lines indicate household growth, and dots indicate social vulnerability; overlaid patterns preserve the combinations.")
    add_body(doc, "A condition is high when it meets the countywide 80th-percentile threshold: GMT +2.5°C precipitation of at least 203.603 mm, projected household growth of at least 746 households, or SVI of at least 0.88386. The legend reports both the thresholds and tract counts.")
    add_section_heading(doc, "6.3", "Configurable two-factor overlap mode")
    add_body(doc, "Enable the overlap switch, select factor A and factor B, and set a countywide percentile threshold from the 60th to the 95th percentile. Each tract is classified as neither high, factor-A-only, factor-B-only, or both high. The dashboard reports the exact numeric cutoffs and summarizes the both-high tracts and their projected residents, households, and jobs.")
    add_section_heading(doc, "6.4", "Tract, ZIP-code, and commissioner-precinct location")
    add_body(doc, "Enter an exact 11-digit GEOID to zoom and select a tract, or click a tract on the map. Enter a five-digit ZIP code to zoom to and outline its boundary; the optional toggle displays all 155 ZIP boundaries. Select Commissioner Precinct 1–4 to zoom to and outline it; a second toggle displays all precinct boundaries. ZIP and precinct geometry support navigation and screening summaries but do not create official geographic aggregates. The selected-tract panel reports CHEI 2050, SVI, population and household change, hotspot class, parcel housing-unit change, precipitation across all four GMT thresholds, and 2020-versus-2050 population/household/employment values.")
    add_section_heading(doc, "6.5", "Quick Decision View and one-page brief")
    add_body(doc, "The four Quick Decision buttons activate and filter the most relevant layer for highest overall exposure, three-factor overlap, emerging adaptation needs, or selected-place conditions. A selected tract produces exact report values; ZIP and precinct reports provide medians, ranges, high-condition shares, and intersecting-tract counts. The browser opens a print-ready letter-size report with Print / Save as PDF and a clear non-aggregation caveat.")
    add_section_heading(doc, "6.6", "Housing-stock point mode and interpretation")
    add_body(doc, "In standalone mode, the map displays the embedded representative sample. In server mode, every pan or zoom triggers a bounded API request for source points in the visible map extent. Use the full-record density grid for countywide comparison and the point layer for local inspection. CHEI and the other modeled, projected, or interpolated indicators are screening inputs; they do not establish causality, engineering criteria, parcel eligibility, or an official risk determination.")

    # Decision-tool illustration.
    new_page(doc, "06B  DECISION TOOL", "From mapped data to a one-page planning brief", "A secondary workflow for users who need concise, decision-oriented information")
    if REPORT_SCREENSHOT.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        shape = p.add_run().add_picture(str(REPORT_SCREENSHOT), width=Inches(7.05))
        shape._inline.docPr.set("title", "Commissioner precinct decision brief")
        shape._inline.docPr.set("descr", "Print-ready one-page screening summary for a selected Harris County commissioner precinct, including a boundary map, exposure indicators, compound-screening shares, decision findings, limitations, and Print or Save as PDF control.")
        cap = doc.add_paragraph("Figure 2. Browser-generated commissioner precinct decision brief; ZIP and precinct outputs summarize intersecting census tracts and are not official aggregates.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.name = "Liberation Sans"
        cap.runs[0].font.size = Pt(8.5)
        cap.runs[0].font.italic = True
        cap.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    add_key_value_table(doc, [
        ("Census tract report", "Exact tract values and the selected tract's compound-hotspot class."),
        ("ZIP / precinct report", "Screening summary across intersecting census tracts using medians, ranges, shares, and tract counts."),
        ("Output", "Browser Print / Save as PDF; no server-side PDF engine is required."),
        ("Interpretation safeguard", "Every ZIP or precinct brief states that the result is not an official geographic aggregate."),
    ])

    # Deployment page.
    new_page(doc, "07  DEPLOYMENT", "Publish the dashboard with the included container files", "The package is deployment-ready but does not create an external account or domain")
    add_section_heading(doc, "7.1", "Build and run with Docker")
    add_code(doc, """
docker build -t climate-housing-exposure-index .
docker run --rm -p 8050:8050 climate-housing-exposure-index
""")
    add_body(doc, "Test the local container at http://127.0.0.1:8050 and confirm that /api/health returns status=ok.")
    add_section_heading(doc, "7.2", "Deploy to a container host")
    add_step(doc, 1, "Create a repository", "For GitHub Pages, upload index.html and .nojekyll at minimum; the complete package also retains source code, documentation, and prepared data. Do not upload the extracted source geodatabase unless archival sharing is intended.")
    add_step(doc, 2, "Create a Docker web service", "Use the included Dockerfile or render.yaml example. Allocate enough image/storage capacity for the processed data and standalone HTML.")
    add_step(doc, 3, "Verify the service", "Open the platform URL, inspect several layers, pan the housing-point map, and check /api/health.")
    add_step(doc, 4, "Publish the access link", "Use the provider-issued HTTPS address or map a custom domain. That address becomes the public dashboard access link.")
    add_subheading(doc, "Operational endpoints")
    add_key_value_table(doc, [
        ("/", "Dashboard HTML"),
        ("/api/health", "Application and source-point readiness check"),
        ("/api/metadata", "Data audit JSON"),
        ("/api/housing-points", "Viewport-filtered point response; maximum response size is capped"),
    ])
    add_body(doc, "A live external URL is not generated by the offline build environment. The standalone HTML is immediately accessible from the delivered package, and the public-deployment configuration is included for an authorized hosting account.")

    # Page 9.
    new_page(doc, "08  QUALITY, TERMS, AND SOURCES", "Validation record and responsible-use requirements", "Use the dashboard as an exploratory evidence layer, not as a sole decision basis")
    add_section_heading(doc, "8.1", "Completed validation")
    add_key_value_table(doc, [
        ("Data integrity", f"{len(inventory)} GDB layers inventoried; {counts['census_tracts']:,} tracts; {counts['zip_codes']:,} ZIP boundaries; {counts['commissioner_precincts']:,} commissioner precincts; {counts['parcels']:,} parcels; sorted full-point arrays verified."),
        ("Interactive behavior", "Clickable legends across tract, parcel, grid, climate-point, housing-point, hotspot, and kriging views; tract and ZIP search; all-ZIP overlay; precinct selection and all-precinct overlay; four Quick Decision actions; and tract/ZIP/precinct report generation tested without JavaScript exceptions."),
        ("Server behavior", "Health endpoint, dashboard response, and capped combined housing-point query tested successfully."),
        ("Visual review", "Overview cards, contained legends, unclipped distributions, ZIP and precinct controls and outlines, Quick Decision panel, one-page letter report, explanatory panels, map, profiles, methods, footer, and terms inspected at desktop resolution."),
    ])
    add_section_heading(doc, "8.2", "Terms of use")
    add_body(doc, "This dashboard is developed to visualize and explore spatial data related to precipitation extremes, housing stocks, population, and land-use projections. The information is intended for research, educational, and informational purposes only.")
    add_body(doc, "The displayed data are derived from multiple sources and may include modeled, estimated, or projected values. While reasonable efforts have been made to support accuracy and reliability, no guarantee is made regarding completeness, accuracy, or timeliness. Users should not rely solely on this dashboard for decision-making, and the creators and affiliated institutions assume no responsibility or liability for errors, omissions, or damages arising from use.")
    add_body(doc, "Unless otherwise noted, the data and visualizations are provided for non-commercial use. Proper attribution should be given when referencing or reproducing the materials. By accessing and using the dashboard, users acknowledge and agree to these terms.")
    add_section_heading(doc, "8.3", "Source organizations and contact")
    sources = [
        ("Texas Tech University Climate Center", "https://www.depts.ttu.edu/csc/"),
        ("H-GAC Regional Growth Forecast", "https://www.h-gac.com/regional-growth-forecast"),
        ("H-GAC Regional Land Use Information System", "https://datalab.h-gac.com/rluis/"),
        ("Harris Central Appraisal District public data", "https://hcad.org/hcad-online-services/pdata/"),
        ("CDC/ATSDR Social Vulnerability Index", "https://www.atsdr.cdc.gov/place-health/php/svi/svi-data-documentation-download.html"),
    ]
    for label, url in sources:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label + ": ")
        r.font.name = "Liberation Sans"
        r.font.size = Pt(9.2)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(NAVY)
        add_hyperlink(p, url, url)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Questions or feedback: Kaifa Lu / CECREH / Kaifa.Lu@ttu.edu")
    r.font.name = "Liberation Sans"
    r.font.size = Pt(10.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)

    # Core properties.
    doc.core_properties.title = "Climate Housing Exposure Index Dashboard Reproducibility Guide"
    doc.core_properties.subject = "Harris County climate, housing, growth, and social vulnerability dashboard"
    doc.core_properties.author = "CECREH"
    doc.core_properties.keywords = "Harris County, CHEI, climate precipitation, housing, SVI, ZIP code, commissioner precinct, interactive legend, decision brief, compound hotspot, dashboard"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
